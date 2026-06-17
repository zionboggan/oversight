#!/usr/bin/env python3
"""
OVERSIGHT v0.2 comprehensive end-to-end test.

Covers:
   1. Identity + keygen
   2. Text watermarking (L1, L2, L3 semantic)
   3. Image DCT watermarking
   4. PDF metadata marks
   5. DOCX metadata marks
   6. Seal + open (single recipient)
  7. Multi-recipient seal fails closed
   8. Policy enforcement (not_after expired)
   9. Policy enforcement (max_opens counter)
  10. Semantic watermark verification (airgap-strip survivor)
  11. Tamper detection
  12. Merkle transparency log correctness
  13. Perceptual hash lookup (fuzzy match)
"""

import io
import sys
import time
import tempfile
import hashlib
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

from oversight_core import (
    ClassicIdentity, Manifest, Recipient, WatermarkRef,
    content_hash, seal, open_sealed, beacon, watermark,
)
from oversight_core import semantic
from oversight_core.container import seal_multi
from oversight_core.policy import PolicyContext, PolicyViolation
from oversight_core.tlog import TransparencyLog


def banner(m): print(f"\n{'=' * 64}\n  {m}\n{'=' * 64}")
def ok(msg): print(f"  [ok] {msg}")
def fail(msg): print(f"  [FAIL] {msg}"); sys.exit(1)


def main():
    banner("1. Identities")
    issuer = ClassicIdentity.generate()
    alice = ClassicIdentity.generate()
    bob = ClassicIdentity.generate()
    carol = ClassicIdentity.generate()
    ok(f"generated 4 identities")

    banner("2. Text watermarking — L1 + L2 + L3")
    text_lines = [f"Supporting paragraph {i}: we begin to show how this is significant and we must help users find answers." for i in range(60)]
    original_text = "\n".join(text_lines)
    mid_zw = watermark.new_mark_id()
    mid_ws = watermark.new_mark_id()
    mid_sem = watermark.new_mark_id()

    # Apply L3 FIRST (rewrites words), then L2 (trailing whitespace),
    # then L1 (zero-width unicode). This order preserves semantic marks
    # even after L1 insertion, because semantic verification strips ZW chars.
    t = semantic.apply_semantic(original_text, mid_sem)
    t = watermark.embed_ws(t, mid_ws)
    t = watermark.embed_zw(t, mid_zw)
    plaintext = t.encode("utf-8")
    ok(f"applied L3/L2/L1 marks; bytes={len(plaintext)}")

    banner("3. Semantic recovery survives airgap-strip")
    # Simulate airgap-strip: remove zero-width, normalize whitespace
    airgap_stripped = t
    for zw in ("\u200b", "\u200c", "\u200d"):
        airgap_stripped = airgap_stripped.replace(zw, "")
    # normalize trailing whitespace
    airgap_stripped = "\n".join(line.rstrip() for line in airgap_stripped.splitlines())

    # Verify L1 and L2 are dead
    l1_survived = len(watermark.extract_zw(airgap_stripped)) > 0
    l2_result = watermark.extract_ws(airgap_stripped)
    l2_survived = l2_result is not None and l2_result == mid_ws
    print(f"  L1 survived airgap-strip: {l1_survived} (expected False)")
    print(f"  L2 survived airgap-strip: {l2_survived} (expected False)")
    if l1_survived or l2_survived:
        fail("L1 or L2 unexpectedly survived airgap-strip — test setup bug")

    # Verify L3 semantic DID survive
    result = semantic.verify_semantic(airgap_stripped, mid_sem)
    print(f"  L3 synonym score:      {result['synonyms_score']:.3f} (match={result['synonyms_match']})")
    print(f"  L3 punctuation hits:   {result['punctuation_hits']}")
    print(f"  L3 overall match:      {result['overall_match']}")
    if not result["overall_match"]:
        fail("L3 semantic watermark failed to survive airgap-strip")
    ok("L3 semantic watermark SURVIVED airgap-strip — attribution possible")

    # Negative: a DIFFERENT mark_id should NOT match
    wrong_result = semantic.verify_semantic(airgap_stripped, watermark.new_mark_id())
    if wrong_result["overall_match"] and wrong_result["synonyms_score"] > 0.65:
        fail(f"random mark_id matched (score={wrong_result['synonyms_score']}) — false positive")
    ok(f"L3 rejects wrong mark_id (score={wrong_result['synonyms_score']:.3f})")

    banner("4. Image DCT watermarking")
    try:
        from PIL import Image
        import numpy as np
        from oversight_core.formats import image as img_fmt
        # Make a test image
        arr = np.random.RandomState(42).randint(64, 200, (256, 256, 3), dtype=np.uint8)
        pil = Image.fromarray(arr)
        buf = io.BytesIO(); pil.save(buf, format="PNG")
        orig_bytes = buf.getvalue()

        img_mark = watermark.new_mark_id()
        marked_bytes = img_fmt.embed(orig_bytes, img_mark, alpha=0.10)
        ok(f"embedded into image: {len(marked_bytes)} bytes")

        match, score = img_fmt.verify(marked_bytes, img_mark)
        print(f"  correct mark score: {score:+.4f} (match={match})")
        if not match:
            fail(f"DCT watermark verify FAILED for correct mark_id")

        wrong_mark = watermark.new_mark_id()
        wrong_match, wrong_score = img_fmt.verify(marked_bytes, wrong_mark)
        print(f"  wrong mark score:   {wrong_score:+.4f} (match={wrong_match})")
        if wrong_match:
            fail("DCT watermark verify matched WRONG mark_id (false positive)")
        ok("image DCT watermark verifies correctly and rejects wrong marks")

        # JPEG recompression attack
        from PIL import Image as _I
        pil2 = _I.open(io.BytesIO(marked_bytes))
        jpeg_buf = io.BytesIO(); pil2.save(jpeg_buf, format="JPEG", quality=75)
        match_after_jpeg, score_after_jpeg = img_fmt.verify(jpeg_buf.getvalue(), img_mark)
        print(f"  post-JPEG-q75 score: {score_after_jpeg:+.4f} (match={match_after_jpeg})")
        if match_after_jpeg:
            ok("image watermark SURVIVED JPEG recompression (q=75)")
        else:
            print(f"  [note] image watermark weakened by JPEG recompression (score below threshold)")

        phash = img_fmt.perceptual_hash(marked_bytes)
        ok(f"perceptual hash: {phash}")
    except Exception as e:
        fail(f"image test error: {e}")

    banner("5. PDF marks")
    try:
        from pypdf import PdfWriter as _PW
        from oversight_core.formats import pdf as pdf_fmt
        # Build a simple PDF using reportlab if available, else skip
        try:
            from reportlab.pdfgen import canvas
            buf = io.BytesIO()
            c = canvas.Canvas(buf)
            c.drawString(100, 750, "Confidential — test document")
            c.save()
            pdf_bytes = buf.getvalue()
        except ImportError:
            # minimal PDF — pypdf can write an empty doc
            w = _PW()
            w.add_blank_page(width=612, height=792)
            buf = io.BytesIO(); w.write(buf)
            pdf_bytes = buf.getvalue()

        pdf_mark = watermark.new_mark_id()
        marked_pdf = pdf_fmt.embed(pdf_bytes, pdf_mark, issuer_id="acme", file_id="pdf-test-1")
        ok(f"embedded into PDF: {len(marked_pdf)} bytes")

        extracted = pdf_fmt.extract(marked_pdf)
        if extracted["mark_id"] != pdf_mark.hex():
            fail(f"PDF mark mismatch: got {extracted['mark_id']}, expected {pdf_mark.hex()}")
        if extracted["issuer_id"] != "acme":
            fail(f"PDF issuer mismatch")
        ok(f"PDF mark recovered: {extracted['mark_id']}")
    except Exception as e:
        fail(f"PDF test error: {e}")

    banner("6. DOCX marks")
    try:
        from docx import Document
        from oversight_core.formats import docx as docx_fmt
        doc = Document()
        doc.add_paragraph("Confidential test document")
        doc.add_paragraph("Second paragraph of content")
        buf = io.BytesIO(); doc.save(buf)
        docx_bytes = buf.getvalue()

        docx_mark = watermark.new_mark_id()
        marked_docx = docx_fmt.embed(docx_bytes, docx_mark, issuer_id="acme", file_id="docx-test-1")
        ok(f"embedded into DOCX: {len(marked_docx)} bytes")

        ext = docx_fmt.extract(marked_docx)
        if ext["mark_id"] != docx_mark.hex():
            fail(f"DOCX mark mismatch: got {ext['mark_id']}, expected {docx_mark.hex()}")
        ok(f"DOCX mark recovered: {ext['mark_id']}")
    except Exception as e:
        fail(f"DOCX test error: {e}")

    banner("7. Single-recipient seal + open (regression)")
    rec = Recipient(recipient_id="alice@corp", x25519_pub=alice.x25519_pub.hex(), ed25519_pub=alice.ed25519_pub.hex())
    m = Manifest.new("test.txt", content_hash(plaintext), len(plaintext), "acme", issuer.ed25519_pub.hex(), rec, "http://localhost:8765", "text/plain")
    m.watermarks = [
        WatermarkRef(layer="L1_zero_width", mark_id=mid_zw.hex()),
        WatermarkRef(layer="L2_whitespace", mark_id=mid_ws.hex()),
        WatermarkRef(layer="L3_semantic", mark_id=mid_sem.hex()),
    ]
    blob = seal(plaintext, m, issuer.ed25519_priv, alice.x25519_pub)
    recovered, mm = open_sealed(blob, alice.x25519_priv)
    if recovered != plaintext:
        fail("recovered plaintext mismatch")
    ok(f"seal/open round-trip OK ({len(blob)} bytes)")

    banner("8. Multi-recipient seal fails closed")
    m2 = Manifest.new("multi.txt", content_hash(plaintext), len(plaintext), "acme", issuer.ed25519_pub.hex(), rec, "http://localhost:8765", "text/plain")
    try:
        seal_multi(
            plaintext, m2, issuer.ed25519_priv,
            [alice.x25519_pub, bob.x25519_pub, carol.x25519_pub],
        )
        fail("seal_multi should be disabled until the manifest can bind multiple recipients")
    except Exception as e:
        ok(f"multi-recipient seal correctly rejected: {type(e).__name__}")

    banner("9. Policy: not_after (expired)")
    expired_m = Manifest.new("exp.txt", content_hash(plaintext), len(plaintext), "acme", issuer.ed25519_pub.hex(), rec, "http://localhost:8765")
    expired_m.policy["not_after"] = int(time.time()) - 60
    expired_blob = seal(plaintext, expired_m, issuer.ed25519_priv, alice.x25519_pub)
    try:
        open_sealed(expired_blob, alice.x25519_priv)
        fail("expired file should NOT open")
    except PolicyViolation as e:
        ok(f"expired file correctly rejected: {e}")

    banner("10. Policy: max_opens counter")
    with tempfile.TemporaryDirectory() as td:
        ctx = PolicyContext(state_dir=Path(td), mode="LOCAL_ONLY", jurisdiction="GLOBAL")
        capped_m = Manifest.new("capped.txt", content_hash(plaintext), len(plaintext), "acme", issuer.ed25519_pub.hex(), rec, "http://localhost:8765")
        capped_m.policy["max_opens"] = 2
        capped_blob = seal(plaintext, capped_m, issuer.ed25519_priv, alice.x25519_pub)

        # First two opens succeed
        for i in range(2):
            pt, _ = open_sealed(capped_blob, alice.x25519_priv, policy_ctx=ctx)
            if pt != plaintext:
                fail(f"open {i+1} recovered wrong plaintext")
        ok("first 2 opens succeeded")

        # Third open should fail
        try:
            open_sealed(capped_blob, alice.x25519_priv, policy_ctx=ctx)
            fail("3rd open should have been rejected")
        except PolicyViolation as e:
            ok(f"3rd open correctly rejected: {e}")

    banner("11. Tamper detection (ciphertext + manifest)")
    bad = bytearray(blob)
    bad[-1] ^= 0x01
    try:
        open_sealed(bytes(bad), alice.x25519_priv)
        fail("ciphertext tamper should have been caught")
    except Exception as e:
        ok(f"ciphertext tamper rejected: {type(e).__name__}")

    bad2 = bytearray(blob)
    bad2[30] ^= 0x01
    try:
        open_sealed(bytes(bad2), alice.x25519_priv)
        fail("manifest tamper should have been caught")
    except Exception as e:
        ok(f"manifest tamper rejected: {type(e).__name__}")

    banner("12. Merkle transparency log")
    with tempfile.TemporaryDirectory() as td:
        reg_key = ClassicIdentity.generate()
        tl = TransparencyLog(td, signing_key_hex=reg_key.ed25519_priv.hex())
        idx0 = tl.append({"event": "test", "i": 0})
        idx1 = tl.append({"event": "test", "i": 1})
        idx2 = tl.append({"event": "test", "i": 2})
        idx3 = tl.append({"event": "test", "i": 3})
        if tl.size() != 4:
            fail(f"tlog size {tl.size()} != 4")
        ok(f"appended 4 entries, size={tl.size()}")

        head = tl.signed_head()
        ok(f"signed head: size={head['size']} root={head['root'][:16]}...")

        # Inclusion proof for index 2
        proof = tl.inclusion_proof(idx2)
        if proof is None:
            fail("inclusion proof for valid index returned None")
        if proof["index"] != idx2:
            fail(f"proof index mismatch")
        ok(f"inclusion proof for idx={idx2}: {len(proof['proof'])} sibling hashes")

        # Adding a new entry changes the root
        root_before = tl.root()
        tl.append({"event": "test", "i": 4})
        if tl.root() == root_before:
            fail("root did not change after append")
        ok("root changes on append (append-only integrity)")

    banner("13. Perceptual hash deterministic")
    try:
        from oversight_core.formats import image as img_fmt
        ph1 = img_fmt.perceptual_hash(marked_bytes)
        ph2 = img_fmt.perceptual_hash(marked_bytes)
        if ph1 != ph2:
            fail("perceptual hash not deterministic")
        ok(f"phash deterministic: {ph1}")
    except Exception as e:
        fail(f"phash test error: {e}")

    banner("ALL TESTS PASSED")


def test_e2e_v2_full_round_trip():
    """Pytest entry point. The scenario is one end-to-end flow with internal
    assertions; pytest's value here is collection + CI integration, not
    per-step granularity."""
    main()


if __name__ == "__main__":
    main()
